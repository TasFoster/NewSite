# -*- coding: utf-8 -*-
"""Генератор шаблона отчёта по Django-проекту (TicketManager-подобному)."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn

GRAY = RGBColor(0x80, 0x80, 0x80)
FONT = "Times New Roman"

doc = Document()

# --- базовый стиль документа ---
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(14)
style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
pf = style.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
pf.space_after = Pt(0)

# --- поля страницы (по ГОСТ) ---
sec = doc.sections[0]
sec.left_margin = Cm(3)
sec.right_margin = Cm(1.5)
sec.top_margin = Cm(2)
sec.bottom_margin = Cm(2)


def set_run(run, *, size=14, bold=False, italic=False, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    return run


def para(text="", *, align=None, bold=False, italic=False, size=14,
         color=None, before=0, after=6, indent=True):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if indent and align in (None, WD_ALIGN_PARAGRAPH.JUSTIFY):
        p.paragraph_format.first_line_indent = Cm(1.25)
    if text:
        set_run(p.add_run(text), size=size, bold=bold, italic=italic, color=color)
    return p


def hint(text):
    """Серая курсивная подсказка-плейсхолдер."""
    return para("[ " + text + " ]", italic=True, color=GRAY,
                align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=10)


def heading(num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    set_run(p.add_run(f"{num}. {text}"), size=14, bold=True)
    return p


def body(text):
    return para(text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=6)


# ============================================================
# ТИТУЛЬНЫЙ ЛИСТ
# ============================================================
def tline(text, *, bold=False, size=14, after=0, color=None):
    return para(text, align=WD_ALIGN_PARAGRAPH.CENTER, bold=bold,
                size=size, after=after, color=color, indent=False)


tline("МИНИСТЕРСТВО ОБРАЗОВАНИЯ И НАУКИ", after=2)
tline("[ наименование учебного заведения ]", color=GRAY, after=2)
tline("[ факультет / кафедра ]", color=GRAY, after=0)

for _ in range(6):
    tline("")

tline("ОТЧЁТ", bold=True, size=18, after=4)
tline("по лабораторной / практической работе", after=2)
tline("на тему:", after=2)
tline("«[ название вашего проекта ]»", bold=True, size=16, after=0)

for _ in range(6):
    tline("")

# блок "Выполнил / Проверил" — выравнивание по правому краю
def right_block(label, value):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run(label + " "), size=14)
    set_run(p.add_run(value), size=14, italic=True, color=GRAY)

right_block("Выполнил(а) студент(ка) группы", "[ группа ]")
right_block("", "[ Фамилия И.О. ]")
para("", after=4)
right_block("Проверил(а):", "[ должность, Фамилия И.О. ]")

for _ in range(6):
    tline("")

tline("[ город ], 2026", color=GRAY)

doc.add_page_break()

# ============================================================
# СОДЕРЖАТЕЛЬНАЯ ЧАСТЬ
# ============================================================

# 1. Цель и задачи
heading(1, "Цель и задачи работы")
body("Цель работы — разработать веб-приложение на фреймворке Django, "
     "реализующее [ кратко суть вашего приложения ].")
para("Для достижения цели поставлены задачи:", align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=2)
for t in ("спроектировать модели данных и связи между ними;",
          "реализовать основные сценарии работы пользователя (CRUD-операции);",
          "настроить маршрутизацию, представления (views) и шаблоны;",
          "[ добавьте свои задачи: регистрация, права доступа и т. д. ]"):
    lp = doc.add_paragraph(style="List Bullet")
    lp.paragraph_format.space_after = Pt(2)
    set_run(lp.add_run(t), size=14,
            italic=t.startswith("["), color=GRAY if t.startswith("[") else None)

# 2. Описание функционала
heading(2, "Описание функционала")
hint("Опишите, что умеет ваше приложение, с точки зрения пользователя: "
     "какие есть страницы и действия, кто и что может делать. "
     "3–6 предложений или короткий список основных возможностей.")
body("Пример (по проекту-образцу TicketManager): приложение позволяет создавать "
     "команды, добавлять в них участников, заводить задачи (тикеты) и управлять "
     "их статусом (Backlog → To Do → In Progress → Review → Done). Каждая задача "
     "привязана к команде и может быть назначена на исполнителя.")

# 3. Структура проекта и модели данных
heading(3, "Структура проекта и модели данных")
hint("Перечислите приложения (apps) вашего проекта и ключевые модели с их "
     "основными полями. Можно вставить фрагмент кода модели или схему связей.")
body("Пример структуры проекта-образца:")
mono = doc.add_paragraph()
mono.paragraph_format.left_indent = Cm(1.25)
mono.paragraph_format.space_after = Pt(6)
code = ("project/\n"
        "  accounts/   — пользователи\n"
        "  teams/      — команды и участники (Team, TeamMembership)\n"
        "  tickets/    — задачи (Ticket)\n"
        "  templates/  — HTML-шаблоны\n"
        "  manage.py")
r = mono.add_run(code)
set_run(r, size=11)
r.font.name = "Consolas"
r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
body("Основные модели: команда (название, капитан, участники), задача "
     "(заголовок, описание, статус, исполнитель, автор, даты создания и "
     "завершения). [ Опишите модели вашего проекта аналогично. ]")

# 4. Использованные технологии  (заполненный раздел-образец)
heading(4, "Использованные технологии")
for t in ("язык программирования — Python 3;",
          "веб-фреймворк — Django 6.0.5;",
          "база данных — SQLite (по умолчанию в Django);",
          "шаблонизатор Django Template Language, HTML/CSS;",
          "[ добавьте, если использовали: Bootstrap, PostgreSQL и др. ]"):
    lp = doc.add_paragraph(style="List Bullet")
    lp.paragraph_format.space_after = Pt(2)
    set_run(lp.add_run(t), size=14,
            italic=t.startswith("["), color=GRAY if t.startswith("[") else None)

# 5. Запуск проекта  (заполненный раздел-образец)
heading(5, "Запуск проекта")
body("Для запуска проекта необходимо выполнить команды:")
cmds = doc.add_paragraph()
cmds.paragraph_format.left_indent = Cm(1.25)
cmds.paragraph_format.space_after = Pt(6)
cmd_text = ("python -m venv .venv\n"
            ".venv\\Scripts\\activate\n"
            "pip install -r requirements.txt\n"
            "python manage.py migrate\n"
            "python manage.py runserver")
rc = cmds.add_run(cmd_text)
set_run(rc, size=11)
rc.font.name = "Consolas"
rc._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
body("После запуска приложение доступно по адресу http://127.0.0.1:8000/.")

# 6. Скриншоты работы приложения
heading(6, "Скриншоты работы приложения")
hint("Вставьте сюда снимки экрана основных страниц приложения и подпишите их. "
     "Под каждым скриншотом — подпись вида «Рисунок 1 — Главная страница».")
para("[ место для скриншота ]", align=WD_ALIGN_PARAGRAPH.CENTER,
     italic=True, color=GRAY, after=2, indent=False)
para("Рисунок 1 — [ название ]", align=WD_ALIGN_PARAGRAPH.CENTER,
     size=12, color=GRAY, after=10, indent=False)
para("[ место для скриншота ]", align=WD_ALIGN_PARAGRAPH.CENTER,
     italic=True, color=GRAY, after=2, indent=False)
para("Рисунок 2 — [ название ]", align=WD_ALIGN_PARAGRAPH.CENTER,
     size=12, color=GRAY, after=10, indent=False)

# 7. Выводы
heading(7, "Выводы")
hint("Кратко (3–5 предложений): что было сделано, какие навыки получены, "
     "с какими трудностями столкнулись и как их решили, что можно улучшить.")

out = "C:\\Users\\User\\Documents\\Projects\\NewSite\\Шаблон_отчёта.docx"
doc.save(out)
print("Saved:", out)
